#!/usr/bin/env python3
"""
Generate Felix Lin's resume in Terry's Base Template v6 layout.

Approach: copy template, then walk each table cell and replace contents,
preserving the original paragraph/run formatting by inspecting the first
run in each cell as a style reference.

Usage:
    python3 build-felix-resume.py [output-path]

Run from any directory; paths are absolute.
"""

from __future__ import annotations
import copy
import os
import shutil
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt
    from copy import deepcopy
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx --break-system-packages")
    sys.exit(1)


# ---------------------------------------------------------------------------
# Source paths
# ---------------------------------------------------------------------------

TEMPLATE = Path(
    "/Users/terrylin/Library/CloudStorage/GoogleDrive-terrylin921@gmail.com/"
    "My Drive/AI/Jobs/Terry_Lin_Base_Template_v6.docx"
)

DEFAULT_OUT = Path(
    "/Users/terrylin/Library/CloudStorage/GoogleDrive-terrylin921@gmail.com/"
    "My Drive/AI/Jobs/Felix Lin Resume v1.docx"
)


# ---------------------------------------------------------------------------
# Felix's resume content
# ---------------------------------------------------------------------------

NAME = "FELIX LIN"
TAGLINE = "Environmental Science · Field Conservation & GIS"

CONTACT_LEFT = "☎  713 459 8357"
CONTACT_MID = "@  linfelix26@gmail.com"
CONTACT_RIGHT = "📍  The Woodlands, TX · open to relocation"


# Right column blocks
SUMMARY_BLOCK = """SUMMARY
Environmental Science graduate (USF, May 2026, GPA 3.8, honors) with hands-on field experience at the National Park Service, sustainable land management with City Grazing, and community organizing with Refuse Refuse. GIS-certified, comfortable with watershed and habitat data, and grounded in years of outdoor stewardship work since high school. Looking to contribute to conservation, cleantech, geothermal, or marine-focused organizations doing real on-the-ground work.

EDUCATION
B.S. Environmental Science, Minor in Computer Science
University of San Francisco · San Francisco, CA · May 2026
GPA 3.8 · Graduated with honors
GIS (Geographic Information System) Certificate

SKILLS
Field surveys & native flora/fauna identification
GIS mapping (ArcGIS / QGIS) — habitat zones, watershed
Coastal watershed data collection & analysis
Sustainable vegetation management (incl. targeted goat grazing)
Community organizing & program coordination
Public-facing education & visitor outreach
Trail systems maintenance
Plant care & landscape stewardship

SOFTWARE/TOOLS/LANGUAGES
GIS · ArcGIS / QGIS
Microsoft Excel · Microsoft Word
Google Workspace (Docs, Sheets, Slides)
Canva (marketing materials)
Mandarin Chinese (fluent)
English (native)

ACHIEVEMENTS
President & Founder — Sustainable Development Club, USF
President's Volunteer Service Award (Gold)
National Merit Scholarship Commended
Member — Environmental Engineering & Science Club, SEEDS (Strategies for Ecology Education, Diversity, and Sustainability), Food Recovery Network
External — The Woodlands GREEN, Stem4Free, Animal Rescue Club

CERTIFICATIONS
GIS (Geographic Information System) — University of San Francisco"""


# Left column — Experience block
# Each job: title line, company/dates line, then bullets
EXPERIENCE_BLOCK = """EXPERIENCE
Field Operations Intern\t📍 Marin Headlands, CA
National Park Service\t📅 10/2023 – 12/2023
Led visitor education programs and maintained extensive trail systems to enhance park accessibility and safety
Conducted comprehensive field surveys of native fauna and flora and assisted with GIS mapping of sensitive habitat zones
Collected and analyzed coastal watershed data to monitor saltwater intrusion patterns and ecosystem dynamics

Sustainable Land Management Intern\t📍 San Francisco, CA
City Grazing\t📅 05/2023 – 07/2023
Implemented sustainable vegetation management using targeted goat grazing for wildfire prevention and control
Monitored and analyzed herd health metrics, grazing patterns, and vegetation impact through detailed field data collection

Community Outreach Intern\t📍 San Francisco, CA
Refuse Refuse\t📅 10/2022 – 12/2022
Planned and led multiple neighborhood cleanups, fostering a safer and cleaner environment across the city
Spearheaded an innovative cigarette disposal solution, designing and assembling portable trays for local distribution
Managed communication channels and produced engaging marketing materials in Excel and Canva

Counter\t📍 The Woodlands, TX
Tide\t📅 11/2021 – 04/2022
Consistently delivered exceptional service to 20-30 daily customers, achieving high tips and positive call-in reviews
Trained and mentored 7 new hires on company processes, enterprise software, and customer service excellence

Bird Specialist\t📍 The Woodlands, TX
Raising Cane's\t📅 06/2021 – 08/2021
Led daily operations and motivated team, setting new service-time records while handling 200+ customers per shift

Assistant Groundskeeper\t📍 The Woodlands, TX
Carlton Woods Turf & Plant Care\t📅 06/2020 – 08/2020
Administered expert plant care and monitoring, maintaining diverse landscape, plant health, aesthetics, and aeration
Gained deeper understanding and appreciation for local flora and fauna, inspiring future environmental pathway

Busser\t📍 The Woodlands, TX
Churrascos Steakhouse\t📅 06/2019 – 08/2019
Provided efficient bussing for 75+ daily customers and pursued additional event catering opportunities"""


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def replace_cell_text(cell, new_text: str) -> None:
    """
    Replace cell contents with new_text, preserving paragraph and run
    formatting from the first run found.

    Each line in `new_text` becomes a paragraph. Tab characters within a line
    become tab characters in the run (template uses tabs for the right-align
    location/date columns).

    A blank line ("") creates an empty paragraph (visual breathing room).
    """
    # Capture style references from the first run we can find
    style_runs = []
    for p in cell.paragraphs:
        for r in p.runs:
            if r.text.strip():
                style_runs.append(r)
                break
        if style_runs:
            break

    # Default style: first non-empty run we found, else None
    style_run = style_runs[0] if style_runs else None

    # Capture all paragraph elements so we can reuse first as template, drop rest
    paragraphs = list(cell.paragraphs)

    # Remove all but the first paragraph
    for p in paragraphs[1:]:
        p._element.getparent().remove(p._element)

    # Clear runs from the first paragraph
    first_p = paragraphs[0]
    for r in list(first_p.runs):
        r._element.getparent().remove(r._element)
    # Also clear any text directly in pPr (rare but possible)

    # Add new content
    lines = new_text.split("\n")

    def add_run(paragraph, text):
        run = paragraph.add_run(text)
        if style_run is not None:
            # Copy basic font properties
            run.font.name = style_run.font.name
            if style_run.font.size:
                run.font.size = style_run.font.size
            if style_run.font.bold is not None:
                run.font.bold = style_run.font.bold
            if style_run.font.italic is not None:
                run.font.italic = style_run.font.italic
        return run

    is_first = True
    for line in lines:
        if is_first:
            p = first_p
            is_first = False
        else:
            p = cell.add_paragraph()
            # Copy paragraph format from first paragraph
            p.paragraph_format.alignment = first_p.paragraph_format.alignment
            if first_p.paragraph_format.left_indent is not None:
                p.paragraph_format.left_indent = first_p.paragraph_format.left_indent

        if not line.strip():
            # Empty paragraph for spacing — skip adding any run
            continue

        # Handle tabs explicitly: split on tab, add runs separated by tab chars
        if "\t" in line:
            parts = line.split("\t")
            for i, part in enumerate(parts):
                if i > 0:
                    # Insert a tab character
                    add_run(p, "\t")
                if part:
                    add_run(p, part)
        else:
            add_run(p, line)


# ---------------------------------------------------------------------------
# Build
# ---------------------------------------------------------------------------

def build(out_path: Path) -> None:
    if not TEMPLATE.exists():
        print(f"ERROR: template not found: {TEMPLATE}")
        sys.exit(1)

    # Translate paths to VM if running in sandbox
    out_path.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy(TEMPLATE, out_path)

    doc = Document(str(out_path))

    # Map sections by inspecting first paragraph of cell
    print(f"Document has {len(doc.tables)} tables")
    for ti, t in enumerate(doc.tables):
        print(f"  Table {ti}: {len(t.rows)} rows x {len(t.columns)} cols")

    # Table 0 — name + tagline cell
    # Original: "TERRY LIN\nPrincipal Product Manager - Growth and Platform"
    if len(doc.tables) >= 1:
        cell = doc.tables[0].rows[0].cells[0]
        replace_cell_text(cell, f"{NAME}\n{TAGLINE}")

    # Table 1 — three contact cells
    if len(doc.tables) >= 2:
        contact_row = doc.tables[1].rows[0]
        if len(contact_row.cells) >= 3:
            replace_cell_text(contact_row.cells[0], CONTACT_LEFT)
            replace_cell_text(contact_row.cells[1], CONTACT_MID)
            replace_cell_text(contact_row.cells[2], CONTACT_RIGHT)

    # Table 2 — left = experience, right = summary/skills/etc
    if len(doc.tables) >= 3:
        body_row = doc.tables[2].rows[0]
        if len(body_row.cells) >= 2:
            replace_cell_text(body_row.cells[0], EXPERIENCE_BLOCK)
            replace_cell_text(body_row.cells[1], SUMMARY_BLOCK)

    doc.save(str(out_path))
    print(f"\n✓ Saved: {out_path}")


if __name__ == "__main__":
    out = Path(sys.argv[1]) if len(sys.argv) > 1 else DEFAULT_OUT
    build(out)
